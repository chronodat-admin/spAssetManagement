#!/usr/bin/env python3
"""Rebuild Asset Management end-user guide (.docx and .md) with screenshots."""

from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "docs" / "Asset-Management-User-Guide.docx"
MD_PATH = ROOT / "docs" / "Asset-Management-User-Guide.md"
GUIDE_IMAGES = ROOT / "docs" / "user-guide" / "images"
ARCH_IMAGES = ROOT / "assets" / "website" / "infographic"
PACKAGE_SCREENSHOT = ROOT / "sharepoint" / "assets" / "screenshot-1.png"
SITE_URL = "https://chronodat.sharepoint.com/sites/ChronodatProdApps/spfx"

sys.path.insert(0, str(ROOT / "scripts"))
from lib.read_app_version import read_app_version  # noqa: E402

VERSION = read_app_version()
GUIDE_DATE = date.today().strftime("%B %Y")

SCREENSHOT_INDEX: list[tuple[str, str]] = [
    ("01-dashboard.png", "Dashboard — KPI cards and charts"),
    ("02-all-assets.png", "All Assets register with search and filters"),
    ("03-assigned-to-me.png", "Assets assigned to the signed-in user"),
    ("04-available-assets.png", "Assets available for assignment"),
    ("05-assign-asset.png", "Assign Asset operation"),
    ("06-return-asset.png", "Return Asset operation"),
    ("07-book-asset.png", "Book Asset for temporary use"),
    ("08-request-asset.png", "Request a new asset"),
    ("09-scan-asset.png", "Scan barcode or QR label"),
    ("10-inventory.png", "Record physical inventory scans"),
    ("11-software-licenses.png", "Software license tracking"),
    ("12-maintenance.png", "Maintenance records"),
    ("13-reports.png", "Report Builder"),
    ("14-depreciation.png", "Depreciation schedules"),
    ("15-audit-log.png", "Audit trail"),
    ("16-categories.png", "Categories lookup list"),
    ("17-vendors.png", "Vendors lookup list"),
    ("18-locations.png", "Locations lookup list"),
    ("19-settings-general.png", "Settings → General"),
    ("20-settings-appearance.png", "Settings → Appearance"),
    ("21-settings-forms.png", "Settings → Forms"),
    ("22-settings-tags.png", "Settings → Tags"),
    ("23-settings-subscription.png", "Settings → Subscription"),
    ("24-settings-roles.png", "Settings → Roles & Permissions"),
]


def add_title(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold


def add_steps(doc: Document, steps: list[str]) -> None:
    for step in steps:
        doc.add_paragraph(step, style="List Number")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_image(doc: Document, path: Path, caption: str, width: float = 5.9) -> None:
    if not path.is_file():
        add_para(doc, f"[Screenshot pending: {path.name} — run npm run docs:user-guide:capture]")
        return
    doc.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def pick(*candidates: Path) -> Path | None:
    for path in candidates:
        if path.is_file():
            return path
    return None


def img(name: str) -> Path:
    return GUIDE_IMAGES / name


def list_guide_images() -> list[Path]:
    if not GUIDE_IMAGES.is_dir():
        return []
    return sorted(GUIDE_IMAGES.glob("*.png"))


def build_markdown() -> None:
    lines: list[str] = [
        "# Asset Management — End User Guide",
        "",
        f"**Version {VERSION}** · {GUIDE_DATE}",
        "",
        "This guide explains how to use Asset Management in Microsoft 365 SharePoint Online "
        "and Microsoft Teams. Screenshots were captured from the live app at:",
        "",
        f"`{SITE_URL}`",
        "",
        "---",
        "",
        "## Screenshot index",
        "",
        "| File | Shows |",
        "|------|-------|",
    ]
    for file_name, desc in SCREENSHOT_INDEX:
        rel = f"user-guide/images/{file_name}"
        lines.append(f"| [{file_name}]({rel}) | {desc} |")

    sections: list[tuple[str, str, list[str], str | None]] = [
        (
            "1. Getting started",
            "1.1 Add the web part and run setup",
            [
                "Open your SharePoint site and edit a modern page.",
                "Click **+** to add a web part and search for **Asset Management**.",
                "Publish the page and open it as a site owner.",
                "Click **Complete Setup** on the banner. Wait until all SharePoint lists are provisioned.",
                "Open **Settings → General** to set the app display name.",
            ],
            "01-dashboard.png",
        ),
        (
            "2. Dashboard",
            "2.1 Review asset KPIs",
            [
                "Open **Dashboard** from the sidebar MAIN section.",
                "Review summary cards for total assets, assigned, available, and overdue items.",
                "Use charts to see distribution by status, category, and location.",
                "Click a card or chart segment to drill into the matching asset list.",
            ],
            "01-dashboard.png",
        ),
        (
            "3. Asset register",
            "3.1 Browse and search assets",
            [
                "Open **All Assets** from the ASSETS section.",
                "Use the search box to find assets by ID, title, serial number, or tag.",
                "Switch between table, list, and card views.",
                "Apply filters for status, category, location, or assigned user.",
                "Click **Create Asset** (or **New Asset**) to add a new item.",
            ],
            "02-all-assets.png",
        ),
        (
            "3. Asset register",
            "3.2 Create or edit an asset",
            [
                "Click **Create Asset** or open an asset title from any list.",
                "On the **General** tab, enter Title, Category, Status, and other required fields.",
                "Complete **Financial**, **Assignment**, and **Maintenance** tabs as needed.",
                "If the category has a linked form template, fill in the extra fields shown below the main tabs.",
                "Add attachments in the attachments section.",
                "Open the **Activity** tab on an existing asset to view SharePoint version history.",
                "Click **Save** to persist changes.",
            ],
            "02-all-assets.png",
        ),
        (
            "3. Asset register",
            "3.3 Filtered asset views",
            [
                "**Assigned To Me** — assets where you are the custodian.",
                "**Available** — assets ready to assign or book.",
                "**In Repair** / **Retired** / **Deleted Assets** — lifecycle-specific views.",
            ],
            "03-assigned-to-me.png",
        ),
        (
            "4. Operations",
            "4.1 Assign an asset",
            [
                "Open **Assign Asset** from the OPERATIONS section.",
                "Search for the asset by ID, title, or scan label.",
                "Select the person to assign to and set optional notes or due date.",
                "Click **Assign** to complete the assignment.",
            ],
            "05-assign-asset.png",
        ),
        (
            "4. Operations",
            "4.2 Return an asset",
            [
                "Open **Return Asset**.",
                "Find the assigned asset and confirm the return condition.",
                "Click **Return** to mark the asset as available again.",
            ],
            "06-return-asset.png",
        ),
        (
            "4. Operations",
            "4.3 Book and request assets",
            [
                "**Book Asset** — reserve an asset for a date range.",
                "**Request Asset** — submit a request when you need equipment.",
                "**My Requests** / **Manage Requests** — track and approve requests.",
            ],
            "07-book-asset.png",
        ),
        (
            "4. Operations",
            "4.4 Scan and inventory",
            [
                "Open **Scan Asset** and enter or scan a barcode/QR label to locate an asset.",
                "Open **Inventory** to record a physical scan during an inventory cycle.",
                "Enter a value in **Scan label** before **Record scan** becomes active.",
            ],
            "10-inventory.png",
        ),
        (
            "5. Licenses and maintenance",
            "5.1 Software licenses",
            [
                "Open **Software Licenses** to track license seats and assignments.",
                "Add licenses with vendor, seat count, and renewal dates.",
                "Link license usage to assets or users as your process requires.",
            ],
            "11-software-licenses.png",
        ),
        (
            "5. Licenses and maintenance",
            "5.2 Maintenance",
            [
                "Open **Maintenance** to log repairs, inspections, and service schedules.",
                "Create maintenance records linked to assets.",
            ],
            "12-maintenance.png",
        ),
        (
            "6. Analysis and reporting",
            "6.1 Report Builder",
            [
                "Open **Reports** from the ANALYSIS section.",
                "Choose a data source (Assets, Vendors, Locations, etc.).",
                "Select columns and optional filters.",
                "Click **Generate Report** to preview, then **Download CSV** to export.",
            ],
            "13-reports.png",
        ),
        (
            "6. Analysis and reporting",
            "6.2 Depreciation and audit",
            [
                "**Depreciation** — review depreciation schedules configured on assets.",
                "**Audit Log** — search create, update, and delete actions across the app.",
            ],
            "14-depreciation.png",
        ),
        (
            "7. Lookup lists",
            "7.1 Master data",
            [
                "Use **Categories**, **Sub-Categories**, **Vendors**, **Locations**, and **Projects** "
                "in the LOOKUPS section to maintain reference data.",
                "Click **Add new**, complete fields, and **Save**.",
                "Lookup values appear in asset forms and filters.",
            ],
            "16-categories.png",
        ),
        (
            "8. Settings (administrators)",
            "8.1 General configuration",
            [
                "Open **Settings** from the ADMIN section (site owners and app administrators only).",
                "**General** — app display name and header links.",
                "**Appearance** — theme, colors, and layout.",
                "**Forms** — field visibility per list (AM_Assets, AM_Vendors, etc.).",
                "**Tags** — colored tags for filtering assets.",
                "**Subscription** — 14-day trial and yearly licensing.",
                "**Roles & Permissions** — assign app roles and UI permissions.",
            ],
            "19-settings-general.png",
        ),
        (
            "9. Tips and troubleshooting",
            "9.1 Common issues",
            [
                "**Setup banner** — run Complete Setup as site owner.",
                "**Settings not visible** — ask an app administrator to add you under App Administrators.",
                "**Record scan disabled** — enter text in Scan label first.",
                "**Email notifications** — tenant admin must approve Microsoft Graph Mail.Send.",
            ],
            None,
        ),
    ]

    for section_title, subsection, steps, screenshot in sections:
        lines.extend(["", f"## {section_title}", "", f"### {subsection}", ""])
        for i, step in enumerate(steps, 1):
            lines.append(f"{i}. {step}")
        if screenshot:
            rel = f"user-guide/images/{screenshot}"
            lines.extend(["", f"![{subsection}]({rel})", ""])

    lines.extend(
        [
            "",
            "---",
            "",
            f"*Generated for Asset Management v{VERSION}. "
            "Refresh screenshots: `npm run docs:user-guide:all`*",
            "",
        ]
    )

    MD_PATH.parent.mkdir(parents=True, exist_ok=True)
    MD_PATH.write_text("\n".join(lines), encoding="utf-8")
    print(f"Wrote {MD_PATH}")


def build_docx() -> None:
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Asset Management")
    run.bold = True
    run.font.size = Pt(24)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("End User Guide — Step-by-Step Instructions").font.size = Pt(14)

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver.add_run(f"Version {VERSION}  •  {GUIDE_DATE}").font.size = Pt(11)

    doc.add_paragraph()
    add_para(
        doc,
        "This guide explains how to deploy, configure, and use Asset Management in "
        "Microsoft 365 SharePoint Online and Microsoft Teams. UI screenshots are captured "
        f"from {SITE_URL} and stored in docs/user-guide/images/.",
    )

    add_title(doc, "Screenshot index", 2)
    add_table(doc, ["File", "Shows"], list(SCREENSHOT_INDEX))

    doc.add_page_break()

    # 1. Getting Started
    add_title(doc, "1. Getting Started")
    add_para(
        doc,
        "Asset Management is a SharePoint Framework (SPFx) web part for tracking hardware, "
        "software licenses, assignments, maintenance, depreciation, and inventory — all within "
        "a single modern interface on your SharePoint site.",
    )

    add_title(doc, "1.1 Requirements", 2)
    add_bullets(
        doc,
        [
            "Microsoft 365 SharePoint Online",
            "Modern browser: Edge, Chrome, Firefox, or Safari",
            "Site owner permission for initial setup and Settings",
            "Recommended: dedicated subsite so asset lists stay isolated",
            "Optional: Exchange Online mailbox for notification emails",
        ],
    )

    add_title(doc, "1.2 Deploy the app (tenant administrator)", 2)
    add_steps(
        doc,
        [
            "Build or obtain sharepoint/solution/asset-management.sppkg (npm run ship).",
            "Upload the .sppkg to the SharePoint App Catalog.",
            "Deploy the package and make it available to sites.",
            "On the target site, confirm the app appears in Site contents.",
        ],
    )

    add_title(doc, "1.3 Add the web part to a page", 2)
    add_steps(
        doc,
        [
            "Create or edit a modern SharePoint page.",
            "Add the Asset Management web part and publish.",
            "Open the page while signed in as a site owner.",
        ],
    )
    add_image(
        doc,
        pick(PACKAGE_SCREENSHOT, img("01-dashboard.png")) or PACKAGE_SCREENSHOT,
        "Figure 1 — Asset Management on a SharePoint page",
    )

    add_title(doc, "1.4 Run first-time setup", 2)
    add_steps(
        doc,
        [
            "Click Complete Setup on the banner when the web part loads.",
            "Wait while lists (AM_Assets, lookups, settings) are created and seed data is applied.",
            "Confirm the banner shows that lists are ready.",
            "Open Settings → General to set the app name.",
            "Open Settings → Appearance to choose theme and colors.",
        ],
    )

    add_title(doc, "1.5 Roles and permissions", 2)
    add_table(
        doc,
        ["Role", "Typical user", "Can do", "Cannot do"],
        [
            [
                "Site owner",
                "SharePoint site collection owner",
                "Run Complete Setup, manage list permissions, open Settings",
                "Approve Graph Mail.Send (tenant admin)",
            ],
            [
                "App administrator",
                "IT asset manager added in Settings",
                "Open Settings, manage lookups, workflows, tags, forms",
                "Override SharePoint list security without site owner",
            ],
            [
                "Asset manager",
                "Operations lead",
                "Create and edit assets, assign, return, run reports",
                "Open Settings unless granted app administrator",
            ],
            [
                "Asset viewer",
                "Staff, auditor",
                "View dashboard, lists, export CSV",
                "Create or edit assets without list permissions",
            ],
        ],
    )
    add_para(
        doc,
        "SharePoint list permissions on AM_Assets and lookup lists control create, edit, and delete. "
        "Settings → Roles & Permissions controls which pages appear in the UI for each app role.",
    )

    doc.add_page_break()

    # 2. Tour
    add_title(doc, "2. Application Tour")
    add_title(doc, "2.1 Top bar", 2)
    add_bullets(
        doc,
        [
            "App name and brand icon (Settings → General)",
            "Procedure link — optional URL to your asset policy document",
            "Dark mode toggle — stored per browser",
            "Create Asset — opens the new asset form",
            "User profile pill when SharePoint chrome is hidden",
        ],
    )

    add_title(doc, "2.2 Sidebar navigation", 2)
    add_table(
        doc,
        ["Section", "Page", "Purpose"],
        [
            ["MAIN", "Dashboard", "KPI cards, charts, recent activity"],
            ["ASSETS", "All Assets", "Full register with search and export"],
            ["ASSETS", "Assigned To Me / Available / In Repair / Retired", "Filtered lifecycle views"],
            ["OPERATIONS", "Assign / Return / Book / Request", "Day-to-day asset workflows"],
            ["OPERATIONS", "Scan Asset / Inventory", "Barcode and physical inventory"],
            ["OPERATIONS", "Software Licenses / Maintenance", "License seats and service records"],
            ["ANALYSIS", "Reports / Depreciation / Audit Log", "Reporting and compliance trail"],
            ["LOOKUPS", "Categories / Vendors / Locations / Projects", "Master data"],
            ["ADMIN", "Settings", "Site owner and app administrator configuration"],
        ],
    )
    add_image(doc, img("01-dashboard.png"), "Figure 2 — Dashboard")

    doc.add_page_break()

    # 3. Assets
    add_title(doc, "3. Managing Assets — Step by Step")
    add_title(doc, "3.1 Browse the register", 2)
    add_steps(
        doc,
        [
            "Open All Assets from the sidebar.",
            "Use search to find assets by ID, title, serial number, or tag.",
            "Switch table, list, or card view.",
            "Filter by status, category, location, or assignee.",
            "Export filtered results to CSV when available.",
        ],
    )
    add_image(doc, img("02-all-assets.png"), "Figure 3 — All Assets")

    add_title(doc, "3.2 Create a new asset", 2)
    add_steps(
        doc,
        [
            "Click Create Asset in the top bar or on the All Assets page.",
            "On the General tab, enter Title, Category, Status, and required fields.",
            "Complete Financial, Assignment, and Maintenance tabs as needed.",
            "If the category has a linked Form Template, complete extra fields below the main tabs.",
            "Add attachments if needed.",
            "Click Save. An Asset ID is assigned from Settings → Numbering.",
        ],
    )

    add_title(doc, "3.3 View, edit, and activity history", 2)
    add_steps(
        doc,
        [
            "Click an asset title to open the detail panel.",
            "Click Edit, change fields, then Save.",
            "Open the Activity tab to view SharePoint version history.",
            "Use row actions → Delete when you have delete permission.",
        ],
    )

    add_title(doc, "3.4 Filtered views", 2)
    add_bullets(
        doc,
        [
            "Assigned To Me — assets where you are custodian",
            "Available — assets ready to assign",
            "In Repair / Retired / Deleted Assets — lifecycle-specific lists",
        ],
    )
    add_image(doc, img("03-assigned-to-me.png"), "Figure 4 — Assigned To Me")
    add_image(doc, img("04-available-assets.png"), "Figure 5 — Available Assets")

    doc.add_page_break()

    # 4. Operations
    add_title(doc, "4. Operations — Step by Step")
    add_title(doc, "4.1 Assign an asset", 2)
    add_steps(
        doc,
        [
            "Open Assign Asset.",
            "Search for the asset by ID, title, or scan label.",
            "Select the assignee and optional due date or notes.",
            "Click Assign.",
        ],
    )
    add_image(doc, img("05-assign-asset.png"), "Figure 6 — Assign Asset")

    add_title(doc, "4.2 Return an asset", 2)
    add_steps(
        doc,
        [
            "Open Return Asset.",
            "Locate the assigned asset.",
            "Confirm condition and click Return.",
        ],
    )
    add_image(doc, img("06-return-asset.png"), "Figure 7 — Return Asset")

    add_title(doc, "4.3 Book and request", 2)
    add_steps(
        doc,
        [
            "Book Asset — reserve an asset for a date range.",
            "Request Asset — submit a request for equipment.",
            "My Requests / Manage Requests — track and approve requests.",
        ],
    )
    add_image(doc, img("07-book-asset.png"), "Figure 8 — Book Asset")
    add_image(doc, img("08-request-asset.png"), "Figure 9 — Request Asset")

    add_title(doc, "4.4 Scan and inventory", 2)
    add_steps(
        doc,
        [
            "Scan Asset — enter or scan a barcode/QR label to open the asset.",
            "Inventory — record a physical scan during an inventory cycle.",
            "On Inventory, enter a Scan label before Record scan is enabled.",
        ],
    )
    add_image(doc, img("09-scan-asset.png"), "Figure 10 — Scan Asset")
    add_image(doc, img("10-inventory.png"), "Figure 11 — Inventory")

    doc.add_page_break()

    # 5. Licenses & Maintenance
    add_title(doc, "5. Software Licenses & Maintenance")
    add_title(doc, "5.1 Software licenses", 2)
    add_steps(
        doc,
        [
            "Open Software Licenses.",
            "Add licenses with vendor, seat count, and renewal date.",
            "Track seat usage against assets or users.",
        ],
    )
    add_image(doc, img("11-software-licenses.png"), "Figure 12 — Software Licenses")

    add_title(doc, "5.2 Maintenance", 2)
    add_steps(
        doc,
        [
            "Open Maintenance.",
            "Create records for repairs, inspections, or scheduled service.",
            "Link maintenance entries to assets.",
        ],
    )
    add_image(doc, img("12-maintenance.png"), "Figure 13 — Maintenance")

    doc.add_page_break()

    # 6. Analysis
    add_title(doc, "6. Reports, Depreciation & Audit")
    add_title(doc, "6.1 Report Builder", 2)
    add_steps(
        doc,
        [
            "Open Reports.",
            "Choose a data source and select columns.",
            "Add optional filters.",
            "Generate Report to preview; Download CSV for full export.",
        ],
    )
    add_image(doc, img("13-reports.png"), "Figure 14 — Report Builder")

    add_title(doc, "6.2 Depreciation", 2)
    add_para(
        doc,
        "Open Depreciation to review schedules for assets with financial data and "
        "depreciation method configured.",
    )
    add_image(doc, img("14-depreciation.png"), "Figure 15 — Depreciation")

    add_title(doc, "6.3 Audit Log", 2)
    add_para(
        doc,
        "Open Audit Log to search create, update, and delete actions. "
        "Administrators can also review Settings → Audit Log.",
    )
    add_image(doc, img("15-audit-log.png"), "Figure 16 — Audit Log")

    doc.add_page_break()

    # 7. Lookups
    add_title(doc, "7. Lookup Lists — Step by Step")
    add_steps(
        doc,
        [
            "Open Categories, Sub-Categories, Vendors, Locations, or Projects from LOOKUPS.",
            "Click Add new and complete visible fields.",
            "Click Save. Values appear in asset forms and filters.",
        ],
    )
    add_image(doc, img("16-categories.png"), "Figure 17 — Categories")
    add_image(doc, img("17-vendors.png"), "Figure 18 — Vendors")
    add_image(doc, img("18-locations.png"), "Figure 19 — Locations")

    doc.add_page_break()

    # 8. Settings
    add_title(doc, "8. Settings Reference (Administrators)")
    add_para(
        doc,
        "Open Admin → Settings. Site owners and app administrators configure the app here.",
    )
    add_table(
        doc,
        ["Setting area", "What you configure"],
        [
            ["General", "App name, procedure document URL"],
            ["Appearance", "Theme, colors, layout, hide SharePoint chrome"],
            ["Dashboard", "Dashboard name and summary card behavior"],
            ["Forms", "Field visibility per entity (AM_Assets, AM_Vendors, etc.)"],
            ["Form Templates", "Category-linked extra fields on asset forms"],
            ["Asset Status", "Custom statuses and workflow labels"],
            ["Subscription", "14-day trial and yearly subscription"],
            ["App Administrators", "Users who can open Settings"],
            ["Roles & Permissions", "App roles and UI page permissions"],
            ["Tags", "Colored tags for asset categorization"],
            ["Numbering", "Auto Asset ID formats"],
            ["Email Integration", "Graph, Power Automate, or Chronodat API delivery"],
            ["Notification Workflows", "Email on create, assign, overdue, etc."],
            ["Audit Log", "Searchable activity log"],
        ],
    )
    add_image(doc, img("19-settings-general.png"), "Figure 20 — Settings → General")
    add_image(doc, img("20-settings-appearance.png"), "Figure 21 — Settings → Appearance")
    add_image(doc, img("21-settings-forms.png"), "Figure 22 — Settings → Forms")
    add_image(doc, img("22-settings-tags.png"), "Figure 23 — Settings → Tags")
    add_image(doc, img("23-settings-subscription.png"), "Figure 24 — Settings → Subscription")
    add_image(doc, img("24-settings-roles.png"), "Figure 25 — Settings → Roles & Permissions")

    doc.add_page_break()

    # 9. Teams
    add_title(doc, "9. SharePoint vs Microsoft Teams")
    add_table(
        doc,
        ["Topic", "SharePoint page", "Teams tab"],
        [
            ["Data storage", "Lists in current SharePoint site", "Same — backing SharePoint site"],
            ["Setup", "Run on hosting site", "Run on backing SharePoint site"],
            ["Settings", "Full access for site owners / app admins", "Same"],
            ["SharePoint chrome hiding", "Available in Appearance", "Disabled in Teams"],
        ],
    )
    add_steps(
        doc,
        [
            "Deploy the .sppkg to the tenant app catalog.",
            "Sync to Teams from the catalog entry.",
            "Add the app as a channel tab or personal app.",
            "Complete setup on the backing SharePoint site if not already done.",
        ],
    )

    doc.add_page_break()

    # 10. Troubleshooting
    add_title(doc, "10. Tips & Troubleshooting")
    add_table(
        doc,
        ["Issue", "What to try"],
        [
            ["Setup banner won't go away", "Complete Setup as site owner"],
            ["Can't open Settings", "Ask an app administrator to add you"],
            ["Record scan disabled", "Enter a Scan label first on Inventory"],
            ["Email not sent", "Approve Graph Mail.Send; check Email Integration settings"],
            ["Create Asset disabled", "Finish setup; confirm Add on AM_Assets list"],
            ["Subscription banner", "Open Settings → Subscription or contact administrator"],
        ],
    )

    add_title(doc, "10.1 Regenerate this guide", 2)
    add_bullets(
        doc,
        [
            f"Set PLAYWRIGHT_BASE_URL to your SharePoint page (example: {SITE_URL}).",
            "Run npm run test:e2e:setup once to create e2e/.auth/user.json.",
            "Run npm run docs:user-guide:all to capture screenshots and rebuild this document.",
        ],
    )

    doc.add_paragraph()
    add_para(
        doc,
        f"Document generated for Asset Management v{VERSION}. "
        "See README.md and docs/ for technical architecture.",
    )

    DOC_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOC_PATH)
    print(f"Wrote {DOC_PATH}")


def build() -> None:
    build_docx()
    build_markdown()
    images = list_guide_images()
    if not images:
        print(f"Warning: no screenshots in {GUIDE_IMAGES} — run npm run docs:user-guide:capture")
    else:
        print(f"Found {len(images)} screenshots in {GUIDE_IMAGES}")


if __name__ == "__main__":
    build()
