#!/usr/bin/env python3
"""Build docs/Asset-Management-End-User-Adoption-Guide.docx for rollout and daily use."""

from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "docs" / "Asset-Management-End-User-Adoption-Guide.docx"
BRAND = RGBColor(0x0D, 0x94, 0x88)
BODY = RGBColor(0x1E, 0x29, 0x3B)
MUTED = RGBColor(0x64, 0x74, 0x8B)

APP_ICON = ROOT / "assets" / "brand" / "app-icon.png"
MARKETING = ROOT / "assets" / "store" / "listing" / "screenshots" / "marketing"
PROCESS_FLOW = ROOT / "assets" / "docs" / "process-flow"

sys.path.insert(0, str(ROOT / "scripts"))
from lib.read_app_version import read_app_version  # noqa: E402

VERSION = read_app_version()
GUIDE_DATE = date.today().strftime("%B %Y")


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = BRAND if level == 1 else BODY


def add_para(doc: Document, text: str, *, bold: bool = False, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_steps(doc: Document, steps: list[str]) -> None:
    for step in steps:
        doc.add_paragraph(step, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], "E6F7F6")
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, value in enumerate(row):
            row_cells[c_idx].text = value
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str, width: float = 6.0) -> None:
    if not path.is_file():
        add_para(doc, f"[Image not found: {path.name}]")
        return
    doc.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = MUTED
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Cover
    if APP_ICON.is_file():
        icon_para = doc.add_paragraph()
        icon_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        icon_para.add_run().add_picture(str(APP_ICON), width=Inches(1.1))
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Asset Management")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BRAND

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("End User Adoption Guide")
    sub_run.font.size = Pt(18)
    sub_run.font.color.rgb = BODY

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_run = tagline.add_run("Get started, build habits, and roll out Asset Management across your organization")
    tag_run.font.size = Pt(12)
    tag_run.font.color.rgb = MUTED

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Version {VERSION}  •  {GUIDE_DATE}")
    meta_run.font.size = Pt(11)
    meta_run.font.color.rgb = MUTED

    doc.add_paragraph()
    intro = doc.add_paragraph(
        "This guide helps everyday users and team champions adopt Asset Management in Microsoft 365. "
        "It focuses on what you need to know to find equipment, request assets, return items on time, "
        "and work confidently inside SharePoint Online or Microsoft Teams — without administrator jargon."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 1. Welcome
    add_heading(doc, "1. Welcome to Asset Management")
    add_para(
        doc,
        "Asset Management replaces scattered spreadsheets and email chains with one place to track "
        "hardware, software licenses, assignments, and inventory. Everything you see in the app is "
        "stored in SharePoint lists on your organization's site — the same data your IT and facilities "
        "teams rely on for audits and reporting.",
    )
    add_heading(doc, "1.1 What you can do", 2)
    add_table(
        doc,
        ["You want to…", "Start here"],
        [
            ["See equipment assigned to me", "Sidebar → Assigned To Me"],
            ["Find something available to borrow", "Sidebar → Available, or Operations → Book Asset"],
            ["Check out a laptop or device", "Operations → Assign Asset (coordinators)"],
            ["Return equipment", "Operations → Return Asset"],
            ["Look up serial numbers or warranty", "All Assets → open the asset panel"],
            ["See portfolio totals and trends", "Dashboard"],
        ],
    )
    add_image(doc, MARKETING / "dashboard-1366x768.png", "Figure 1 — Dashboard overview")

    add_heading(doc, "1.2 Why adopt this app?", 2)
    add_bullets(
        doc,
        [
            "Single source of truth — one register instead of multiple team spreadsheets.",
            "Self-service — find what is available and what is assigned to you without emailing IT.",
            "Accountability — assignments and returns are logged with dates and notes.",
            "Visibility — managers see utilization, depreciation, and audit history in one place.",
            "Works where you work — open the same app in SharePoint or as a Microsoft Teams tab.",
        ],
    )

    doc.add_page_break()

    # 2. Access
    add_heading(doc, "2. How to Access the App")
    add_para(
        doc,
        "Your organization installs Asset Management once at the tenant level. You do not install anything "
        "on your PC. Ask your site owner or IT contact for the link to your Asset Management page.",
    )
    add_heading(doc, "2.1 Open from SharePoint", 2)
    add_steps(
        doc,
        [
            "Open the SharePoint site or page your organization shared (for example a dedicated Asset Management subsite).",
            "Sign in with your Microsoft 365 work account if prompted.",
            "The Asset Management web part loads the dashboard automatically when setup is complete.",
        ],
    )
    add_heading(doc, "2.2 Open from Microsoft Teams", 2)
    add_steps(
        doc,
        [
            "Open the Teams channel or tab where Asset Management was pinned.",
            "Sign in with the same Microsoft 365 account you use for SharePoint.",
            "Use the sidebar and top bar exactly as in SharePoint — the experience is the same.",
        ],
    )
    add_heading(doc, "2.3 First visit checklist", 2)
    add_table(
        doc,
        ["Check", "What to expect"],
        [
            ["Page loads without errors", "If you see a setup banner, ask a site owner to run Complete Setup first."],
            ["Dashboard appears", "Summary cards and charts show after lists are provisioned."],
            ["You can open Assigned To Me", "If empty, nothing is checked out to you yet — that is normal."],
            ["Create Asset button visible", "If missing, your SharePoint permissions may be view-only; contact your site owner."],
        ],
    )

    doc.add_page_break()

    # 3. Tour
    add_heading(doc, "3. Tour of the Application")
    add_image(doc, MARKETING / "all-features-1366x768.png", "Figure 2 — Application layout and main areas")

    add_heading(doc, "3.1 Top bar", 2)
    add_bullets(
        doc,
        [
            "App name — your organization may customize this in Settings.",
            "Dark mode — toggle light or dark theme; your choice is remembered in this browser.",
            "Portfolio filters — narrow dashboard and list views by business or project when your site uses them.",
            "Create Asset — add a new item to the register (when you have permission).",
            "User profile — shown when SharePoint chrome is hidden for a cleaner embedded experience.",
        ],
    )

    add_heading(doc, "3.2 Sidebar navigation", 2)
    add_table(
        doc,
        ["Section", "Pages", "When to use"],
        [
            ["Main", "Dashboard", "Start of day — KPIs, charts, recent activity"],
            ["Assets", "All Assets, Assigned To Me, Available, In Repair, Retired, Deleted", "Browse and search the register"],
            ["Operations", "Assign, Return, Book, Booking Details, Software Licenses, Inventory", "Check out, return, reserve, or scan assets"],
            ["Analysis", "Reports, Depreciation, Audit Log", "Exports, finance views, change history"],
            ["Lookups", "Categories, Vendors, Locations, Projects", "Reference data (usually coordinators and admins)"],
            ["Admin", "Settings", "Site owners and app administrators only"],
        ],
    )
    add_image(doc, MARKETING / "feature-grid-1366x768.png", "Figure 3 — Feature areas at a glance")

    doc.add_page_break()

    # 4. First 5 minutes
    add_heading(doc, "4. Your First Five Minutes")
    add_steps(
        doc,
        [
            "Open the Dashboard and scan summary cards for totals and alerts.",
            "Click Assigned To Me — confirm whether anything is checked out to you.",
            "Open Available — browse equipment your organization marks as ready to assign.",
            "Click any asset title to open the detail panel (serial number, location, warranty, notes).",
            "Bookmark the SharePoint page or Teams tab so you can return in one click.",
        ],
    )
    add_para(
        doc,
        "Tip: If you only need to return something today, go straight to Operations → Return Asset.",
        bold=True,
    )

    doc.add_page_break()

    # 5. Common tasks
    add_heading(doc, "5. Common Tasks — Step by Step")

    add_heading(doc, "5.1 View assets assigned to me", 2)
    add_steps(
        doc,
        [
            "In the sidebar, under ASSETS, click Assigned To Me.",
            "Review the list — each row is hardware or equipment assigned to your account.",
            "Click a title to open details: serial number, location, assigned date, and notes.",
            "When you no longer need the item, use Operations → Return Asset (see section 5.4).",
        ],
    )

    add_heading(doc, "5.2 Find available equipment", 2)
    add_steps(
        doc,
        [
            "Click Available in the sidebar.",
            "Use search or column filters to narrow by category, location, or asset type.",
            "Open an asset to confirm specifications before requesting it.",
            "Follow your organization's process: some teams use Book Asset first; others ask a coordinator to Assign Asset.",
        ],
    )

    add_heading(doc, "5.3 Book an asset (temporary reservation)", 2)
    add_steps(
        doc,
        [
            "Go to Operations → Book Asset.",
            "Select an available asset from the list.",
            "Choose the requester (usually yourself) and optional expected return date.",
            "Add notes if needed (project name, event, reason).",
            "Submit — the reservation appears under Booking Details until fulfilled or cancelled.",
        ],
    )
    add_para(
        doc,
        "Booking records intent; physical checkout may still require a coordinator to run Assign Asset.",
    )

    add_heading(doc, "5.4 Return an asset", 2)
    add_steps(
        doc,
        [
            "Go to Operations → Return Asset.",
            "Select the asset you are returning from the dropdown (only assigned items appear).",
            "Add optional notes (condition, accessories returned, issues).",
            "Click Submit — the asset returns to Available status and the assignment is logged.",
        ],
    )
    add_para(doc, "Returning on time keeps inventory accurate and helps colleagues find equipment faster.")

    add_heading(doc, "5.5 Assign an asset (coordinators / IT)", 2)
    add_steps(
        doc,
        [
            "Go to Operations → Assign Asset.",
            "Pick an Available asset and the person receiving it.",
            "Add notes if required by your policy.",
            "Submit — the asset status becomes Assigned and the assignee sees it under Assigned To Me.",
        ],
    )

    add_heading(doc, "5.6 Create or update an asset record", 2)
    add_steps(
        doc,
        [
            "Click Create Asset in the top bar, or open an existing asset and choose Edit.",
            "Complete the form — fields may change by category (laptop, monitor, license, etc.).",
            "Save — the asset appears in All Assets and contributes to dashboard metrics.",
        ],
    )

    add_image(
        doc,
        PROCESS_FLOW / "05-operations-and-analysis.png",
        "Figure 4 — Assign, book, return, and related operations",
    )

    doc.add_page_break()

    # 6. Statuses
    add_heading(doc, "6. Understanding Asset Status")
    add_para(
        doc,
        "Each asset has a status that tells you whether it can be assigned, is in use, or is out of service. "
        "Your organization may customize status labels in Settings; the meanings below are typical.",
    )
    add_table(
        doc,
        ["Status", "Meaning", "What you should do"],
        [
            ["Available", "In stock and ready", "Can be assigned or booked"],
            ["Assigned", "Checked out to someone", "Return when finished; do not reassign without coordinator"],
            ["In Repair", "Being serviced", "Do not assign — check back later or pick another item"],
            ["Retired", "End of life", "Reference only — not for checkout"],
            ["Deleted", "Soft-deleted archive", "Admins only — hidden from normal lists"],
        ],
    )
    add_image(
        doc,
        PROCESS_FLOW / "04-asset-lifecycle-states.png",
        "Figure 5 — Typical asset lifecycle",
    )

    doc.add_page_break()

    # 7. Analysis for power users
    add_heading(doc, "7. Reports and Insights (Optional)")
    add_para(
        doc,
        "Most end users stay on Dashboard and Operations pages. Team leads and analysts may also use:",
    )
    add_bullets(
        doc,
        [
            "Reports — build custom exports from asset fields.",
            "Depreciation — finance and asset accounting views.",
            "Audit Log — who changed what and when.",
            "Software Licenses — seat counts and expiry dates.",
            "Inventory — results of physical scan events.",
        ],
    )
    add_image(doc, MARKETING / "analysis-1366x768.png", "Figure 6 — Analysis and reporting area")

    doc.add_page_break()

    # 8. Adoption rollout
    add_heading(doc, "8. Organization Rollout Checklist")
    add_para(
        doc,
        "Site owners and champions can use this checklist to drive adoption. Share this document with "
        "new users during onboarding.",
    )
    add_heading(doc, "8.1 Before launch", 2)
    add_steps(
        doc,
        [
            "Site owner completes setup and verifies Dashboard loads.",
            "Publish the Asset Management page and share the link (SharePoint + optional Teams tab).",
            "Define who may Create Asset, Assign Asset, and edit lookups.",
            "Load initial asset data (import CSV or enter key items manually).",
            "Communicate the return policy and expected status updates.",
        ],
    )
    add_heading(doc, "8.2 Launch week", 2)
    add_steps(
        doc,
        [
            "Send a short announcement with the page link and this guide.",
            "Ask each team lead to verify Assigned To Me for their staff.",
            "Run a 15-minute walkthrough: Dashboard → Assigned To Me → Return Asset.",
            "Designate one coordinator per department for Assign and Book requests.",
        ],
    )
    add_heading(doc, "8.3 Steady state (ongoing)", 2)
    add_bullets(
        doc,
        [
            "Weekly: coordinators review Available vs Assigned counts on the Dashboard.",
            "Monthly: run a report or inventory scan to reconcile physical stock.",
            "Quarterly: retire obsolete assets and update locations/vendors.",
            "Always: return assets the same day they are no longer needed.",
        ],
    )

    add_heading(doc, "8.4 Roles at a glance", 2)
    add_table(
        doc,
        ["Role", "Typical user", "Primary tasks"],
        [
            ["End user", "Any employee", "View Assigned To Me, book/return, read asset details"],
            ["Coordinator", "IT, office manager, team lead", "Assign, book, create assets, run inventory"],
            ["Analyst", "Finance, audit", "Reports, depreciation, audit log"],
            ["Site owner / app admin", "SharePoint owner", "Setup, Settings, permissions, subscription"],
        ],
    )

    doc.add_page_break()

    # 9. Tips
    add_heading(doc, "9. Tips for Success")
    add_bullets(
        doc,
        [
            "Bookmark the page — reduce friction so returns actually happen.",
            "Use notes on assign and return — future you (and auditors) will thank you.",
            "Check Available before buying new equipment — reuse saves budget.",
            "Keep serial numbers accurate — they matter for warranty and Intune-matched devices.",
            "Use Assigned To Me in onboarding — new hires confirm what they received on day one.",
            "Report damaged items promptly — coordinators can set status to In Repair.",
            "Do not edit Settings unless you are an administrator.",
        ],
    )

    add_heading(doc, "9.1 Frequently asked questions", 2)
    add_table(
        doc,
        ["Question", "Answer"],
        [
            ["I do not see Create Asset", "Your SharePoint permissions may be read-only. Ask the site owner."],
            ["The app asks me to subscribe", "Your organization's trial may have ended. Contact an app administrator in Settings."],
            ["I returned hardware but still see it assigned", "Refresh the page. If it persists, contact your coordinator — they can process Return Asset."],
            ["Can I use this on mobile?", "Yes — open the SharePoint page or Teams tab in your mobile browser or Teams app."],
            ["Where is my data stored?", "In SharePoint lists on your organization's site — not in a separate external database."],
            ["Who do I contact for access?", "Your SharePoint site owner or IT service desk."],
        ],
    )

    doc.add_page_break()

    # 10. Quick reference
    add_heading(doc, "10. Quick Reference Card")
    add_table(
        doc,
        ["Goal", "Navigation path"],
        [
            ["My equipment", "Assets → Assigned To Me"],
            ["What's free to use", "Assets → Available"],
            ["Reserve for later", "Operations → Book Asset"],
            ["Hand back device", "Operations → Return Asset"],
            ["Check out to someone", "Operations → Assign Asset"],
            ["Active reservations", "Operations → Booking Details"],
            ["Add new hardware", "Top bar → Create Asset"],
            ["Big-picture metrics", "Main → Dashboard"],
        ],
    )

    add_para(doc, "")
    closing = doc.add_paragraph(
        "Thank you for adopting Asset Management. Accurate assignments and timely returns help your "
        "whole organization work smarter with the equipment you already own."
    )
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if closing.runs:
        closing.runs[0].italic = True
        closing.runs[0].font.color.rgb = MUTED

    DOC_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DOC_PATH))
    print(f"Wrote {DOC_PATH}")


if __name__ == "__main__":
    build()
