"""Generate Asset Management user guide Word document."""
from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Asset-Management-User-Guide.docx"

sys.path.insert(0, str(ROOT / "scripts"))
from lib.read_app_version import read_app_version  # noqa: E402

APP_VERSION = read_app_version()
BRAND = RGBColor(0x0D, 0x94, 0x88)


def set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = BRAND if level == 1 else RGBColor(0x1E, 0x29, 0x3B)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], "E6F7F6")
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, value in enumerate(row):
            row_cells[c_idx].text = value
    doc.add_paragraph()


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Asset Management")
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = BRAND

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("User Guide — Features & How to Use")
    sub_run.font.size = Pt(16)
    sub_run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Version {APP_VERSION}  •  {date.today().strftime('%B %Y')}")
    meta_run.font.size = Pt(11)
    meta_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    doc.add_paragraph()
    intro = doc.add_paragraph(
        "This guide explains how to use the Asset Management application in Microsoft 365 "
        "SharePoint Online and Microsoft Teams. It covers day-to-day risk management, compliance "
        "assessments, reporting, and administrator configuration."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # 1. Overview
    add_heading(doc, "1. Getting Started")
    doc.add_paragraph(
        "Asset Management is a SharePoint Framework (SPFx) application that provides a modern "
        "risk register, analytics dashboard, compliance framework tracking, and configurable "
        "workflows—all within a single web part."
    )
    add_heading(doc, "Requirements", 2)
    add_bullets(
        doc,
        [
            "Microsoft 365 SharePoint Online (not on-premises SharePoint Server)",
            "Modern browser (Edge, Chrome, Firefox, or Safari)",
            "Site owner permission to run initial setup and open Settings",
            "Recommended: dedicated subsite (e.g. /sites/YourSite/AssetManagement) for isolated lists",
        ],
    )
    add_heading(doc, "First-time setup", 2)
    add_numbered(
        doc,
        [
            "Add the Asset Management web part to a modern SharePoint page (full-width recommended).",
            "Sign in as a site owner.",
            "Click Complete Setup on the banner (or open Settings → Run setup).",
            "Wait while the app creates 13 SharePoint lists, seeds lookup data, and registers the risk form.",
            "Configure the app name, appearance, and workflows under Settings as needed.",
        ],
    )
    doc.add_paragraph(
        "After setup, all users with appropriate list permissions can create and manage risks. "
        "Only site owners see the Settings area in the sidebar."
    )

    # 2. App layout
    add_heading(doc, "2. Tour of the Application")
    add_heading(doc, "Top bar", 2)
    add_bullets(
        doc,
        [
            "App name and branding (configured in Settings → General)",
            "Procedure link — opens your risk management procedure document (optional URL in Settings)",
            "Dark mode toggle",
            "Portfolio Filters — filter by business and project across dashboard and risk views",
            "Create New Risk (or New Assessment on compliance pages)",
            "User profile pill — shown when SharePoint top bar is hidden (Settings → Appearance)",
        ],
    )
    add_heading(doc, "Sidebar navigation", 2)
    add_table(
        doc,
        ["Section", "Page", "Purpose"],
        [
            ["Main", "Dashboard", "Portfolio summary, heat map, charts, latest risks"],
            ["Risks", "All Risks", "Full register with search, filters, and export"],
            ["Risks", "Assigned To Me", "Risks where you are the owner/assignee"],
            ["Risks", "Open / In Progress / Closed", "Filtered by workflow status bucket"],
            ["Risks", "Overdue / Due Today / Due This Week", "Date-driven risk views"],
            ["Analysis", "Risk Rating", "Inherent and residual risk matrices"],
            ["Analysis", "Report Builder", "Custom reports and CSV export"],
            ["Compliance", "Compliance Dashboard", "Compliance KPIs and charts"],
            ["Compliance", "Compliance Frameworks", "Framework library and assessments"],
            ["Lookups", "Business", "Business unit master data"],
            ["Lookups", "Projects", "Project portfolio linked to businesses"],
            ["Admin", "Settings", "Site owner configuration (owners only)"],
        ],
    )

    # 3. Risk management
    add_heading(doc, "3. Managing Risks")
    add_heading(doc, "Creating a risk", 2)
    add_numbered(
        doc,
        [
            "Click + Create New Risk in the top bar.",
            "Complete the General tab: title, description, status, category, business, and project.",
            "On the Assessment tab, set likelihood and impact; the app calculates matrix priority.",
            "On the Action tab, add mitigation plan, assign owners, and set due dates.",
            "Add attachments if needed, then click Save.",
            "A Risk ID (e.g. Risk-001) is assigned automatically based on numbering settings.",
        ],
    )
    add_heading(doc, "Viewing and editing", 2)
    doc.add_paragraph(
        "Click any risk title in a list, dashboard table, or heat map to open the detail panel. "
        "Use Edit to modify fields. The Activity tab shows SharePoint version history for saved risks."
    )
    add_heading(doc, "Risk list features", 2)
    add_bullets(
        doc,
        [
            "Switch between table, list, and card views",
            "Search by Risk ID, title, or description",
            "Filter by status and matrix priority (Critical, Major, Moderate, Low)",
            "Apply portfolio filters for business and project",
            "Export filtered results to CSV",
            "Bulk delete (when you have delete permission)",
        ],
    )
    add_heading(doc, "Risk form tabs", 2)
    add_table(
        doc,
        ["Tab", "Typical content"],
        [
            ["General", "Title, description, status, category, business, project, profile type"],
            ["Assessment", "Likelihood, impact, residual ratings, causes, controls, potential cost"],
            ["Action", "Response strategy, mitigation plan, owners, key dates"],
            ["Activity", "Version history timeline (existing risks only)"],
        ],
    )

    # 4. Dashboard & analysis
    add_heading(doc, "4. Dashboard & Risk Analysis")
    add_heading(doc, "Dashboard", 2)
    add_bullets(
        doc,
        [
            "Summary cards: Open, In Progress, Closed, Critical count, average risk age",
            "Inherent risk heat map — click a cell to drill down to matching risks",
            "Severity, status, and priority charts",
            "Latest risks table with tabs (Latest, Assigned to me, Overdue, etc.)",
            "Financial exposure card (optional — enable in Settings → Dashboard)",
        ],
    )
    add_heading(doc, "Risk Rating page", 2)
    doc.add_paragraph(
        "Shows inherent and residual risk matrices side by side. Residual ratings use explicit "
        "potential likelihood/impact when set; otherwise they are derived from control effectiveness. "
        "Click matrix cells or risk names to open details."
    )
    add_heading(doc, "Understanding priority", 2)
    doc.add_paragraph(
        "Matrix priority (Critical, Major, Moderate, Low) is calculated from Likelihood × Impact "
        "using your configured scales (Settings → Likelihood / Consequence). This is separate from "
        "Custom Priorities in Settings, which are organizational labels stored for future use."
    )

    # 5. Compliance
    add_heading(doc, "5. Compliance Management")
    add_heading(doc, "Frameworks", 2)
    doc.add_paragraph(
        "Built-in frameworks (ISO 27001, NIST CSF, GDPR, HIPAA, PCI-DSS, SOC 2, and others) ship "
        "with pre-defined controls. Enable or disable them in Settings → Compliance. Custom frameworks "
        "can be created for organization-specific requirements."
    )
    add_heading(doc, "Assessments workflow", 2)
    add_numbered(
        doc,
        [
            "Go to Compliance Frameworks and click New Assessment.",
            "Enter a name, select an active framework, and optionally set a due date.",
            "Open the assessment and review each control.",
            "Set status (Compliant, Non-Compliant, Partially Compliant, Not Applicable), evidence, and notes.",
            "Monitor progress on the Compliance Dashboard.",
            "Change assessment status to Complete when all controls are assessed.",
        ],
    )
    add_heading(doc, "Compliance Dashboard", 2)
    add_bullets(
        doc,
        [
            "Active frameworks count and assessments in progress",
            "Overall compliance percentage and controls assessed",
            "Charts by assessment and status",
            "Framework cards and recent assessments table",
        ],
    )

    # 6. Lookups
    add_heading(doc, "6. Business & Projects")
    doc.add_paragraph(
        "Business units and projects are master data lists used to classify risks and drive portfolio "
        "filters on the dashboard."
    )
    add_bullets(
        doc,
        [
            "Business — title, industry, region, criticality, owner",
            "Projects — title, code, linked business, status, priority, type, project manager",
            "Create, edit, and delete records from the sidebar lookup pages (permissions apply)",
            "Changes update dashboard filters and risk forms automatically",
        ],
    )

    # 7. Report builder
    add_heading(doc, "7. Report Builder")
    add_numbered(
        doc,
        [
            "Open Analysis → Report Builder.",
            "Choose a data source: Risks, Business, or Projects.",
            "Select columns to include.",
            "Add optional filters (field, operator, value).",
            "Click Generate Report to preview (up to 200 rows).",
            "Click Download CSV for a full export.",
        ],
    )

    # 8. Settings
    add_heading(doc, "8. Settings Reference (Site Owners)")
    doc.add_paragraph(
        "Settings are organized in the sidebar under General, Preferences, and Lookups. "
        "Most pages use a Save settings button at the bottom; lookup list pages save independently."
    )
    add_table(
        doc,
        ["Setting area", "What you configure"],
        [
            ["General", "App name, procedure document URL"],
            ["Appearance", "Theme, colors, layout, hide SharePoint chrome"],
            ["Dashboard", "Dashboard name, dynamic naming, heat-map drill-down, financial card"],
            ["Forms", "Field visibility, labels, and required flags per entity"],
            ["Form Templates", "Extra fields shown when a risk category matches a template"],
            ["Risk Status & Priority", "Custom statuses (with workflow buckets) and priority labels"],
            ["Compliance", "Enable/disable frameworks, seed built-ins, custom frameworks"],
            ["Numbering", "Auto Risk ID and project code formats"],
            ["Tags", "Colored tags for risk categorization"],
            ["Notification Workflows", "Email templates per event (created, assigned, overdue, etc.)"],
            ["Workflow Rules", "Trigger/condition/action rules (requires Power Automate for execution)"],
            ["Email Templates", "Reusable HTML notification templates with variables"],
            ["Scheduled Reports", "Report schedules (requires automation to deliver emails)"],
            ["Audit Log", "Searchable log of create/update/delete actions"],
            ["Lookup lists", "Categories, sub-categories, likelihood, consequence, profile, response, strategy"],
        ],
    )
    add_heading(doc, "Custom statuses explained", 2)
    doc.add_paragraph(
        "Each custom status has a name, color, and Counts as bucket (Open, In Progress, Mitigation, "
        "Resolved, or Closed). Status names sync to the SharePoint Riskstatus field when you save settings. "
        "Sidebar views (Open Risks, In Progress, Closed) group risks by these buckets."
    )

    # 9. Teams
    add_heading(doc, "9. SharePoint vs Microsoft Teams")
    add_table(
        doc,
        ["Topic", "SharePoint page", "Teams tab"],
        [
            ["Data storage", "Lists in current SharePoint site", "Same — backing SharePoint site"],
            ["Setup", "Run on hosting site", "Run on backing SharePoint site"],
            ["Settings", "Full access for site owners", "Same; configure in app Settings"],
            ["SharePoint chrome hiding", "Available in Appearance settings", "Disabled in Teams"],
            ["Native list form customizer", "Works on Risks list URLs", "Not available in Teams"],
        ],
    )

    # 10. Tips & troubleshooting
    add_heading(doc, "10. Tips & Troubleshooting")
    add_table(
        doc,
        ["Issue", "What to try"],
        [
            ["Setup banner won't go away", "Complete Setup as site owner; check Manage Lists permission"],
            ["Can't open Settings", "Only site owners see Settings; ask an owner to grant access"],
            ["Create Risk disabled", "Finish setup; confirm you have Add permission on Risks list"],
            ["Status not in dropdown", "Save Settings → Risk Status & Priority; statuses sync to SharePoint"],
            ["Email notifications not sent", "Configure tenant outbound mail; wire Power Automate to workflow settings"],
            ["Compliance dashboard slow first visit", "First load seeds frameworks/controls; later visits are faster"],
            ["Changes not visible after save", "Lists refresh automatically; try navigating away and back"],
        ],
    )

    add_heading(doc, "Quick reference — keyboard & navigation", 2)
    add_bullets(
        doc,
        [
            "Use sidebar collapse on desktop to maximize content area",
            "Portfolio filters persist per site URL in your browser",
            "Back to top — scroll to footer on long pages",
            "Breadcrumb Home link returns to Dashboard",
        ],
    )

    # Footer note
    doc.add_paragraph()
    footer = doc.add_paragraph(
        f"Document generated for Asset Management v{APP_VERSION}. "
        "For deployment and technical details, see README.md in the solution package."
    )
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
