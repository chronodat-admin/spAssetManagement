#!/usr/bin/env python3
"""Build docs/Asset-Management-Technical-Architecture.docx with infographic diagrams."""

from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "docs" / "Asset-Management-Technical-Architecture.docx"
INFOGRAPHIC = ROOT / "assets" / "website" / "infographic"

sys.path.insert(0, str(ROOT / "scripts"))
from lib.read_app_version import read_app_version  # noqa: E402

VERSION = read_app_version()
GUIDE_DATE = date.today().strftime("%B %Y")

DIAGRAMS: list[tuple[str, str, str]] = [
    (
        "architecture-overview.png",
        "Architecture Overview",
        "One SPFx package deployed across Microsoft 365. The web part and form customizer "
        "run in SharePoint Online and Microsoft Teams; data remains in site-scoped SharePoint lists.",
    ),
    (
        "architecture-platform.png",
        "Platform Stack",
        "Technology layers from Microsoft 365 hosting through SPFx 1.21, React 17, Fluent UI v9, "
        "SharePoint REST, Microsoft Graph Mail.Send, and optional Chronodat external services.",
    ),
    (
        "architecture-components.png",
        "Component Architecture",
        "Presentation components, SPFx web parts and extensions, shared services/contexts, "
        "and packaging configuration in a single asset-management.sppkg bundle.",
    ),
    (
        "architecture-data-flow.png",
        "Data Flow",
        "User interactions flow through the SPFx host and React UI into service classes that "
        "persist to SharePoint lists, call Graph for email, and query the Chronodat Subscription API.",
    ),
    (
        "architecture-modules.png",
        "Application Modules",
        "Risk Register, Compliance, and Administration modules share platform services "
        "(REST, provisioning, permissions, notifications, licensing).",
    ),
    (
        "architecture-surfaces.png",
        "Deployment Surfaces",
        "The same package runs on SharePoint modern pages, Teams channel tabs, Teams personal apps, "
        "and native SharePoint Risks list forms (form customizer).",
    ),
    (
        "architecture-deployment.png",
        "Deployment Flow",
        "From npm run ship through tenant App Catalog, Mail.Send approval, web part placement, "
        "and per-site Complete Setup provisioning.",
    ),
    (
        "architecture-notifications.png",
        "Notification Architecture",
        "Risk lifecycle events trigger NotificationService, which tries Chronodat API first "
        "and falls back to Microsoft Graph sendMail. Failures never block risk saves.",
    ),
    (
        "architecture-security.png",
        "Security & Permissions",
        "SharePoint list permissions, app administrator roles, delegated Graph Mail.Send, "
        "and the optional client-side licensing boundary.",
    ),
    (
        "architecture-compliance.png",
        "Compliance Module",
        "Frameworks, controls, assessments, and dashboard views stored in SharePoint lists "
        "and integrated with the risk register.",
    ),
]


def _diagram_bullets(filename: str) -> list[str]:
    details: dict[str, list[str]] = {
        "architecture-overview.png": [
            "Hub: AssetManagementWebPart + AssetFormCustomizer in one SPFx package",
            "Surfaces: SharePoint pages, Teams tabs, native list forms, Subscription API",
            "Data plane: SharePoint lists, Graph Mail.Send, Chronodat billing API",
        ],
        "architecture-data-flow.png": [
            "All CRUD and analytics run client-side in React services",
            "NotificationService tries Chronodat API before Graph fallback",
            "Setup wizard provisions 13+ lists once per site",
        ],
        "architecture-surfaces.png": [
            "Same .sppkg from tenant App Catalog",
            "Teams personal app and channel tab use SharePoint-backed lists",
            "Form customizer applies to SharePoint Risks list URLs only",
        ],
        "architecture-security.png": [
            "No custom authentication server in the bundle",
            "Mail.Send requires tenant admin approval once per tenant",
            "Licensing gate is client-side; SharePoint data remains accessible via native APIs",
        ],
    }
    return details.get(
        filename,
        ["See companion user guide for operational workflows using this module."],
    )


def add_title(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str, width: float = 6.0) -> None:
    if not path.is_file():
        add_para(doc, f"[Diagram not found: {path.name}. Run npm run assets:infographic after generating raw PNGs.]")
        return
    doc.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
    doc.add_paragraph()


def build() -> None:
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Asset Management")
    run.bold = True
    run.font.size = Pt(24)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Technical Architecture Reference").font.size = Pt(14)

    ver = doc.add_paragraph()
    ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver.add_run(f"Version {VERSION}  •  {GUIDE_DATE}").font.size = Pt(11)

    doc.add_paragraph()
    add_para(
        doc,
        "This document describes the solution architecture for Asset Management — a SharePoint "
        "Framework (SPFx) application for risk registers, compliance tracking, and GRC workflows in "
        "Microsoft 365. Diagrams are AI-generated infographics with the official Chronodat brand; "
        "companion flat SVG sources live in assets/website/ and docs/architecture/.",
    )

    add_title(doc, "Document scope", 2)
    add_bullets(
        doc,
        [
            "Solution topology, data flow, and component map",
            "Deployment surfaces and provisioning flow",
            "Notifications, security, compliance, and licensing boundaries",
            "Technology stack and external service integration",
        ],
    )

    add_title(doc, "Executive summary", 2)
    add_table(
        doc,
        ["Topic", "Design decision"],
        [
            ["Hosting", "SPFx 1.21 web part + form customizer in tenant App Catalog"],
            ["Data residency", "All app data in SharePoint lists on the host web (13+ lists per site)"],
            ["Identity", "Microsoft 365 / SharePoint — no custom auth server in the bundle"],
            ["Email", "Delegated Microsoft Graph Mail.Send; optional Chronodat API fallback"],
            ["Licensing", "Per-site subscription via Chronodat API + Stripe (client-side UI gate)"],
            ["Teams", "Same web part as channel tab and personal app; SharePoint-backed data"],
        ],
    )

    doc.add_page_break()

    add_title(doc, "1. Architecture diagrams")
    add_para(
        doc,
        "The following sections embed high-level infographic diagrams. Regenerate images with "
        "assets/website/infographic/PROMPTS.md and npm run assets:infographic.",
    )

    for idx, (filename, heading, description) in enumerate(DIAGRAMS, start=1):
        add_title(doc, f"1.{idx} {heading}", 2)
        add_para(doc, description)
        add_bullets(
            doc,
            _diagram_bullets(filename),
        )
        add_image(doc, INFOGRAPHIC / filename, f"Figure {idx} — {heading}")
        if idx in (3, 6, 8):
            doc.add_page_break()

    doc.add_page_break()

    add_title(doc, "2. SPFx solution structure")
    add_table(
        doc,
        ["Artifact", "Purpose"],
        [
            ["AssetManagementWebPart", "Main React shell — dashboard, risks, compliance, settings"],
            ["AssetFormCustomizer", "Replaces native SharePoint Risks list new/edit forms"],
            ["TeamsTab / TeamsPersonalApp", "Teams manifest entries pointing to the same web part"],
            ["asset-management.sppkg", "Tenant App Catalog package (skipFeatureDeployment)"],
            ["ListProvisioningService", "Complete Setup — creates lists, seeds lookups, registers forms"],
        ],
    )

    add_title(doc, "2.1 SharePoint lists (provisioned per site)", 2)
    add_bullets(
        doc,
        [
            "Risks, Risk Categories, Risk Sub-Categories, Business, Projects",
            "ComplianceFrameworks, ComplianceControls, ComplianceAssessments, AssessmentEvidence",
            "AppSettings, Administrators, AuditLog, Licenses (schema compatibility)",
            "Lookup lists: likelihood, consequence, profile types, response/strategy, tags, workflows",
        ],
    )

    add_title(doc, "2.2 Key services", 2)
    add_table(
        doc,
        ["Service", "Responsibility"],
        [
            ["RiskService", "CRUD, lookups, app settings via SharePoint REST"],
            ["ComplianceService", "Framework seeding, assessments, control mapping"],
            ["NotificationService", "Workflow emails — Chronodat API then Graph fallback"],
            ["SubscriptionService", "Status, checkout, portal URLs via Chronodat API"],
            ["ListProvisioningService", "Setup wizard, list schema, form customizer registration"],
        ],
    )

    doc.add_page_break()

    add_title(doc, "3. Licensing & subscription integration")
    add_para(
        doc,
        "Licensing is per SharePoint site. The web part calls the Chronodat Subscription API; "
        "Stripe handles billing. Secrets never ship in the SPFx package.",
    )
    add_table(
        doc,
        ["Topic", "Behavior"],
        [
            ["Model", "One licence per SharePoint site; 14-day trial on first status check"],
            ["Product slug", "asset-management-hub"],
            ["Default API", "https://subscription.chronodat.com (web part property override)"],
            ["Access gate", "hasAccess: false shows paywall; SharePoint data remains via native APIs"],
            ["Resilience", "Cached status with 7-day grace window on API connectivity failure"],
            ["Dev bypass", "skipSubscriptionCheck web part property (development only)"],
        ],
    )
    add_bullets(
        doc,
        [
            "GET /api/subscription/status — trial/subscription state",
            "POST /api/subscription/checkout — Stripe Checkout session",
            "POST /api/subscription/portal — Stripe Customer Portal",
            "POST /api/notifications/send — optional hosted email delivery",
        ],
    )
    add_para(
        doc,
        "See docs/architecture/licensing-and-stripe-integration.md for API contracts, "
        "threat boundaries, and hardening roadmap.",
    )

    add_title(doc, "4. Deployment checklist", 2)
    add_bullets(
        doc,
        [
            "Build: npm run clean && npm run ship → sharepoint/solution/asset-management.sppkg",
            "Upload .sppkg to tenant App Catalog and deploy to all sites",
            "Approve Microsoft Graph Mail.Send (SharePoint Admin Center → API access)",
            "Add web part to a modern page or sync to Teams",
            "Run Complete Setup as site owner (provisions lists, seeds data, registers form customizer)",
            "Configure Settings → General, Appearance, Notification Workflows as needed",
            "Optional: set subscriptionApiUrl and verify licensing for production",
        ],
    )

    add_title(doc, "5. Regenerate this document", 2)
    add_bullets(
        doc,
        [
            "Infographic PNGs: assets/website/infographic/ (npm run assets:infographic)",
            "Flat SVG diagrams: npm run assets:architecture",
            "Rebuild Word doc: npm run docs:architecture",
        ],
    )

    doc.add_paragraph()
    add_para(
        doc,
        f"Document generated for Asset Management v{VERSION}. "
        "Chronodat — Asset Management for Microsoft 365.",
    )

    DOC_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(DOC_PATH)
    print(f"Wrote {DOC_PATH}")

    missing = [name for name, _, _ in DIAGRAMS if not (INFOGRAPHIC / name).is_file()]
    if missing:
        print(f"Warning: missing infographics in {INFOGRAPHIC}: {', '.join(missing)}")
    else:
        print(f"Embedded {len(DIAGRAMS)} infographic diagrams from {INFOGRAPHIC}")


if __name__ == "__main__":
    build()
